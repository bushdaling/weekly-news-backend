#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
每周要闻收集工具 - 方案B（全自动后端）v2
重点抓取：新华网、人民网、云南网、迪庆州政府

运行: cd "/Users/a1/Desktop/公文排版工具" && python3 news-collector.py
打开: http://localhost:5000
"""

import io
import re
import json
import base64
import socket
import urllib.request
import urllib.parse
from datetime import datetime, timedelta
from flask import Flask, request, send_file, jsonify
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

app = Flask(__name__)
PORT = 5000

# ====== 全局配置 ======
socket.setdefaulttimeout(8)
HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
    'Accept-Language': 'zh-CN,zh;q=0.9',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
}

# ====== 新闻来源配置 ======
# RSS源
RSS_SOURCES = {
    '中央': [
        # 观察者网 - 时政（实时更新）
        'https://www.guancha.cn/rss/politics.xml',
    ],
    '云南省委': [
        'https://webapp.yunnan.cn/ynwrss/ynw.xml',  # 云南网
    ],
    '迪庆州委': [],  # 迪庆用HTML，见下方
}

# HTML网页源（实时更新）
HTML_SOURCES = {
    '中央': [
        {'url': 'https://www.guancha.cn/politics/', 'source': '观察者网-时政'},
    ],
    '云南省委': [
        {'url': 'https://www.yunnan.cn/', 'source': '云南网'},
    ],
    '迪庆州委': [
        {'url': 'https://www.diqing.gov.cn/xwzx/dqyw.html', 'source': '迪庆州政府-要闻'},
        {'url': 'https://www.diqing.gov.cn/', 'source': '迪庆州政府-首页'},
    ]
}

# 党建/政务关键词
PARTY_KWS = [
    '书记', '党委', '党支部', '党建', '党组', '党风廉政',
    '纪委', '监察', '巡视', '反腐', '廉政', '廉洁', '警示', '通报',
    '习近平', '党中央', '国务院', '中办', '中纪委', '中央军委', '中央政治局',
    '工作会', '部署', '贯彻', '落实', '传达', '学习', '专题',
    '重要指示', '重要讲话', '重要批示', '重要讲',
    '深化改革', '高质量发展', '中国式现代化',
    '从严治党', '主题教育', '纪律', '作风',
    '干部', '人事', '任免', '任职',
    '文件', '条例', '规定', '意见', '通知', '公报',
    '审议', '通过', '决定', '决议',
    '云南省委', '州市', '州委', '州政府',
    '迪庆', '香格里拉', '维西', '德钦',
]

# ====== 网络请求 ======
def fetch(url, is_rss=False):
    try:
        req = urllib.request.Request(url, headers=HEADERS)
        with urllib.request.urlopen(req, timeout=8) as resp:
            ct = resp.headers.get('Content-Type', '')
            enc = 'utf-8'
            if 'gb' in ct.lower(): enc = 'gbk'
            return resp.read().decode(enc, errors='ignore')
    except Exception as e:
        print(f"  ✗ {url[:50]}... {str(e)[:30]}")
        return ''

# ====== RSS解析 ======
def parse_rss(xml, source):
    items = []
    for m in re.finditer(r'<item>(.*?)</item>', xml, re.DOTALL):
        item = m.group(1)
        # 标题
        t = re.search(r'<title><!\[CDATA\[([^\]]+)\]\]></title>', item)
        if not t: t = re.search(r'<title>([^<]+)</title>', item)
        # 链接
        l = re.search(r'<link>([^<]+)</link>', item)
        # 日期
        d = re.search(r'<pubDate>([^<]+)</pubDate>', item)
        # 来源
        s = re.search(r'<source><!\[CDATA\[([^\]]+)\]\]></source>', item) or \
            re.search(r'<source>([^<]+)</source>', item)
        
        if t:
            title = t.group(1).strip()
            if len(title) < 6: continue
            items.append({
                'title': title,
                'url': l.group(1).strip() if l else '',
                'date': d.group(1).strip()[:10] if d else '',
                'source': s.group(1).strip() if s else source,
            })
    return items

# ====== HTML解析（云南网） ======
def parse_yunnan(html, source):
    """解析云南网"""
    items = []
    seen = set()
    patterns = [
        r'<a[^>]+href="(https?://[^"]+yunnan\.cn[^"]+)"[^>]*>\s*([^<]{8,80})\s*</a>',
        r'<a[^>]+href="(/[^"]+)"[^>]+class="[^"]*title[^"]*"[^>]*>([^<]{8,80})</a>',
    ]
    for pat in patterns:
        for url, title in re.findall(pat, html, re.DOTALL):
            title = re.sub(r'<[^>]+>', '', title).strip()
            title = re.sub(r'\s+', ' ', title)
            if len(title) > 6 and title not in seen:
                seen.add(title)
                full_url = url if url.startswith('http') else f'https://www.yunnan.cn{url}'
                items.append({'title': title, 'url': full_url, 'date': '', 'source': source})
    return items[:30]

# ====== HTML解析（备用 - 其他政府网站） ======
def parse_diqing(html, source):
    """通用HTML新闻解析"""
    items = []
    seen = set()
    patterns = [
        r'<a[^>]+href="(https?://[^"]+)"[^>]+class="[^"]*(?:title|news|f)[^"]*"[^>]*>([^<]{8,80})</a>',
        r'<a[^>]+href="(https?://[^"]+)"[^>]*>\s*([^<]{8,80})\s*</a>',
    ]
    for pat in patterns:
        for url, title in re.findall(pat, html, re.DOTALL):
            title = re.sub(r'<[^>]+>', '', title).strip()
            title = re.sub(r'\s+', ' ', title)
            if len(title) > 8 and title not in seen:
                seen.add(title)
                items.append({'title': title, 'url': url, 'date': '', 'source': source})
    return items[:20]

# ====== HTML解析（迪庆政府网站） ======
def parse_gov_html(html, source, base_url='https://www.diqing.gov.cn'):
    """解析迪庆政府网站"""
    items = []
    seen = set()
    
    # 迪庆政府新闻列表页
    patterns = [
        # 标准新闻列表
        r'<a[^>]+href="(/xwzx/[^"]+\.html)"[^>]*>\s*([^<]{8,80})\s*</a>',
        # 完整URL
        r'<a[^>]+href="(https?://www\.diqing\.gov\.cn[^"\s]+\.html)"[^>]*>\s*([^<]{8,80})\s*</a>',
    ]
    
    for pat in patterns:
        for url, title in re.findall(pat, html, re.DOTALL):
            title = re.sub(r'<[^>]+>', '', title).strip()
            title = re.sub(r'\s+', ' ', title)
            if len(title) > 6 and title not in seen:
                seen.add(title)
                if url.startswith('/'):
                    full_url = base_url + url
                else:
                    full_url = url
                date_in_url = re.search(r'/(\d{8}|[a-z0-9]+)\.html', url)
                date_str = ''
                if date_in_url:
                    d = re.search(r'(\d{8})', url)
                    if d:
                        d_str = d.group(1)
                        date_str = f"{d_str[:4]}-{d_str[4:6]}-{d_str[6:8]}"
                items.append({'title': title, 'url': full_url, 'date': date_str, 'source': source})
    
    return items[:30]


# ====== HTML解析（观察者网） ======
def parse_guancha(html, source):
    """解析观察者网"""
    items = []
    seen = set()
    
    patterns = [
        # 观察者网文章格式: /politics/2026_04_02_数字.shtm
        r'<a[^>]+href="(/politics/\d{4}_\d{2}_\d{2}_\d+\.shtml)"[^>]*>\s*([^<]{8,80})\s*</a>',
        r'<a[^>]+href="(/politics/\d{4}_\d{2}_\d{2}_\d+\.htm)"[^>]*>\s*([^<]{8,80})\s*</a>',
    ]
    
    for pat in patterns:
        for url, title in re.findall(pat, html, re.DOTALL):
            title = re.sub(r'<[^>]+>', '', title).strip()
            title = re.sub(r'\s+', ' ', title)
            if len(title) > 6 and title not in seen:
                seen.add(title)
                full_url = 'https://www.guancha.cn' + url
                # 提取日期
                date_in_url = re.search(r'/(\d{4}_\d{2}_\d{2})_', url)
                date_str = ''
                if date_in_url:
                    date_str = date_in_url.group(1).replace('_', '-')
                items.append({'title': title, 'url': full_url, 'date': date_str, 'source': source})
    
    return items[:50]
    items = []
    seen = set()
    
    # 迪庆政府新闻列表页
    patterns = [
        # 标准新闻列表
        r'<a[^>]+href="(/xwzx/[^"]+\.html)"[^>]*>\s*([^<]{8,80})\s*</a>',
        # 带日期的新闻
        r'<a[^>]+href="(/[^"]+\.html)"[^>]*>\s*([^<]{8,80})\s*</a>',
        # 其他政府网站通用
        r'<a[^>]+href="(https?://www\.diqing\.gov\.cn[^"]+\.html)"[^>]*>\s*([^<]{8,80})\s*</a>',
    ]
    
    for pat in patterns:
        for url, title in re.findall(pat, html, re.DOTALL):
            title = re.sub(r'<[^>]+>', '', title).strip()
            title = re.sub(r'\s+', ' ', title)
            if len(title) > 6 and title not in seen:
                seen.add(title)
                # 拼接完整URL
                if url.startswith('/'):
                    full_url = base_url + url
                elif url.startswith('http'):
                    full_url = url
                else:
                    continue
                # 提取日期（URL中通常有）
                date_match = re.search(r'/(\d{4})(\d{2})/(\d{8}|[a-z0-9]+)\.html', url)
                date_str = ''
                if date_match:
                    # 尝试从URL提取日期
                    date_str = date_match.group(1) + '-' + date_match.group(2) + '-01'
                    # 更精确：找日期
                    date_in_url = re.search(r'(\d{8})', url)
                    if date_in_url:
                        d = date_in_url.group(1)
                        date_str = f"{d[:4]}-{d[4:6]}-{d[6:8]}"
                
                items.append({
                    'title': title,
                    'url': full_url,
                    'date': date_str,
                    'source': source,
                })
    
    return items[:30]

# ====== 日期过滤 ======
def date_filter(item, start_date, end_date):
    d = item.get('date', '')
    if not d: return True
    try:
        dt = datetime.strptime(d[:10], '%Y-%m-%d')
        return start_date <= dt <= end_date
    except:
        return True

# ====== 核心抓取函数 ======
def fetch_news(level, start_date, end_date, extra_kws=''):
    results = []
    extra_list = [k.strip() for k in extra_kws.split(',') if k.strip()]
    
    # 1. RSS源
    for url in RSS_SOURCES.get(level, []):
        site = re.search(r'://([^/]+)', url).group(1)
        site = site.replace('www.', '').replace('.com', '')
        print(f"  → {site} RSS...")
        xml = fetch(url, is_rss=True)
        if not xml or len(xml) < 100: continue
        
        items = parse_rss(xml, site)
        print(f"    获取 {len(items)} 条")
        
        for item in items:
            # 日期过滤
            if not date_filter(item, start_date, end_date): continue
            # 关键词过滤（额外）
            if extra_list:
                if not any(kw in item['title'] for kw in extra_list): continue
            # 标记党建类
            item['level'] = level
            item['party'] = any(kw in item['title'] for kw in PARTY_KWS)
            results.append(item)
    
    # 2. HTML网页源
    for cfg in HTML_SOURCES.get(level, []):
        url = cfg['url']
        source = cfg.get('source', url)
        print(f"  → {source}...")
        html = fetch(url)
        if not html: continue
        
        if 'guancha' in url:
            items = parse_guancha(html, source)
        elif 'diqing' in url or 'gov' in url:
            items = parse_gov_html(html, source, base_url='https://www.diqing.gov.cn')
        elif 'yunnan' in url:
            items = parse_yunnan(html, source)
        else:
            items = parse_diqing(html, source)
        
        print(f"    获取 {len(items)} 条")
        for item in items:
            if not date_filter(item, start_date, end_date): continue
            if extra_list:
                if not any(kw in item['title'] for kw in extra_list): continue
            item['level'] = level
            item['party'] = any(kw in item['title'] for kw in PARTY_KWS)
            results.append(item)
    
    # 去重 + 党建优先
    seen, unique = set(), []
    for r in results:
        if r['title'] not in seen:
            seen.add(r['title'])
            unique.append(r)
    unique.sort(key=lambda x: (0 if x.get('party') else 1, x.get('date', '')), reverse=True)
    return unique

# ====== API ======
@app.route('/')
def index():
    with open('/Users/a1/Desktop/公文排版工具/news-collector.html', 'r', encoding='utf-8') as f:
        return f.read()

@app.route('/api/news', methods=['POST'])
def api_news():
    data = request.json
    levels = data.get('levels', ['中央', '云南省委', '迪庆州委'])
    start_str = data.get('start', datetime.now().strftime('%Y-%m-%d'))
    end_str = data.get('end', datetime.now().strftime('%Y-%m-%d'))
    keywords = data.get('keywords', '')
    
    start_date = datetime.strptime(start_str, '%Y-%m-%d')
    end_date = datetime.strptime(end_str, '%Y-%m-%d')
    
    all_news = []
    for level in levels:
        print(f"\n📰 抓取 {level}...")
        news = fetch_news(level, start_date, end_date, keywords)
        all_news.extend(news)
        print(f"  ✓ {level}: {len(news)} 条")
    
    return jsonify({'success': True, 'news': all_news, 'total': len(all_news)})

@app.route('/api/generate', methods=['POST'])
def api_generate():
    data = request.json
    selected = data.get('selected', [])
    if not selected:
        return jsonify({'success': False, 'error': '请先选择要闻'})
    
    excel_buf = make_excel(selected)
    word_buf = make_word(selected)
    
    return jsonify({
        'success': True,
        'files': {
            'excel': base64.b64encode(excel_buf.getvalue()).decode(),
            'word': base64.b64encode(word_buf.getvalue()).decode(),
        }
    })

# ====== 文件生成 ======
def make_excel(news_list):
    wb = Workbook()
    ws = wb.active
    ws.title = "本周要闻"
    
    hf = Font(name='微软雅黑', size=12, bold=True, color='FFFFFF')
    hfill = PatternFill(start_color='C41230', end_color='C41230', fill_type='solid')
    tf = Font(name='微软雅黑', size=11)
    lf = Font(name='微软雅黑', size=11, color='C41230', underline='single')
    ca = Alignment(horizontal='center', vertical='center')
    la = Alignment(horizontal='left', vertical='center', wrap_text=True)
    thin = Side(style='thin', color='DDDDDD')
    bd = Border(left=thin, top=thin, right=thin, bottom=thin)
    
    headers = ['序号', '级别', '新闻标题', '来源', '日期', '链接']
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = hf; c.fill = hfill; c.alignment = ca; c.border = bd
    
    for i, item in enumerate(news_list, 1):
        row = [i, item.get('level',''), item.get('title',''), item.get('source',''), item.get('date',''), item.get('url','')]
        for col, val in enumerate(row, 1):
            c = ws.cell(row=i+1, column=col, value=val)
            c.border = bd; c.alignment = ca if col==1 else la
            if col==3: c.font = tf
            elif col==6 and val:
                c.hyperlink = val; c.font = lf
    
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 14
    ws.column_dimensions['C'].width = 55
    ws.column_dimensions['D'].width = 16
    ws.column_dimensions['E'].width = 14
    ws.column_dimensions['F'].width = 50
    
    buf = io.BytesIO()
    wb.save(buf); buf.seek(0)
    return buf

def make_word(news_list):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.8); sec.right_margin = Cm(2.1)
    
    def sf(run, name='仿宋', size=16, bold=False):
        run.font.name = name; run.font.size = Pt(size)
        run.font.bold = bold; run.font.color.rgb = RGBColor(0,0,0)
        r = run._element; rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {}); rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), name)
    
    def para(text, size=16, bold=False, indent=True, center=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text); sf(run, '仿宋', size, bold)
    
    def heading(text, size=22):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = Pt(32)
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text); sf(run, '方正小标宋简体', size, True)
    
    def section_title(text):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text); sf(run, '黑体', 16, True)
    
    heading("A镇老干部学习资料", 26)
    heading("——每周要闻汇编", 22)
    para(f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}", size=14, bold=False, indent=False)
    doc.add_paragraph()
    
    level_map = [
        ('中央', '一、中央要闻'),
        ('云南省委', '二、云南省委要闻'),
        ('迪庆州委', '三、迪庆州委要闻'),
    ]
    for key, title in level_map:
        items = [n for n in news_list if key in n.get('level', '')]
        if not items: continue
        section_title(title)
        for i, n in enumerate(items, 1):
            para(f"（{i}）{n.get('title', '')}", size=16)
            meta = f"来源：{n.get('source', '')}"
            if n.get('date'): meta += f"　　日期：{n['date']}"
            para(meta, size=14, bold=False, indent=False)
            doc.add_paragraph()
    
    footer = sec.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    run = p.add_run("A镇老干部服务站 编制")
    sf(run, '仿宋', 10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf

# ====== 启动 ======
if __name__ == '__main__':
    print(f"""
╔══════════════════════════════════════════════════╗
║       📰 每周要闻收集工具 v2                    ║
║   重点源：新华网、人民网、云南网、迪庆州政府    ║
╠══════════════════════════════════════════════════╣
║  浏览器打开: http://localhost:{PORT}                ║
║  停止服务: Ctrl+C                               ║
╚══════════════════════════════════════════════════╝
    """)
    app.run(host='0.0.0.0', port=PORT, debug=False, threaded=True)
